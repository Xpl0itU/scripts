#!/usr/bin/env python3

import os
import argparse
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn


def set_styles(doc):
    title_style = doc.styles["Title"]
    title_style.element.rPr.rFonts.set(qn("w:asciiTheme"), "Liberation Sans")
    font = title_style.font
    font.name = "Liberation Sans"
    font.size = Pt(28)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    heading_style = doc.styles["Heading 1"]
    font = heading_style.font
    heading_style.element.rPr.rFonts.set(qn("w:asciiTheme"), "Liberation Sans")
    font.name = "Liberation Sans"
    font.size = Pt(18)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    normal_style = doc.styles["Normal"]
    font = normal_style.font
    font.name = "Liberation Serif"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def create_document():
    doc = Document()

    current_section = doc.sections[-1]
    # Set format to A4
    current_section.page_height = Mm(297)
    current_section.page_width = Mm(210)
    # Rotate to landscape
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height

    return doc


def add_title(doc, title):
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()


def add_toc(doc):
    toc_heading = doc.add_heading("Table Of Contents", level=0)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\f \\o "1-9" \\h'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click to update field."

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

    doc.add_paragraph()
    doc.add_page_break()


def add_borders(doc):
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for border_name in ("left", "right"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "4")
        border_el.set(qn("w:space"), "24")
        border_el.set(qn("w:color"), "auto")
        pg_borders.append(border_el)
    for border_name in ("top", "bottom"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "4")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "auto")
        pg_borders.append(border_el)
    sec_pr.append(pg_borders)


def add_images_to_document(doc, screenshot_folder):
    image_files = [
        f
        for f in os.listdir(screenshot_folder)
        if f.endswith((".png", ".jpg", ".jpeg", ".gif"))
    ]

    for image_file in image_files:
        image_path = os.path.join(screenshot_folder, image_file)
        image_heading = doc.add_heading(image_file, level=1)
        image_heading.font = "Liberation Serif"
        doc.add_picture(image_path, width=Inches(6))
        doc.add_page_break()


def format_sections(doc):
    for section in doc.sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)


def main(title, output_docx, screenshot_folder):
    if not os.path.exists(screenshot_folder):
        print(f"Error: The screenshot folder '{screenshot_folder}' does not exist.")
        return

    doc = create_document()
    set_styles(doc)
    add_title(doc, title)
    add_toc(doc)
    add_images_to_document(doc, screenshot_folder)
    format_sections(doc)
    add_borders(doc)
    doc.save(output_docx)
    print(
        f"Created {output_docx} with {os.path.basename(screenshot_folder)} as the screenshot folder."
    )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Create a Word document with images.")
    parser.add_argument(
        "-t",
        "--title",
        help="Title of the document (default: Screenshots)",
        default="Screenshots",
    )
    parser.add_argument(
        "-o",
        "--output-docx",
        default="screenshots.docx",
        help="Output document filename (default: screenshots.docx)",
    )
    parser.add_argument("screenshot_folder", help="Folder containing image files")
    args = parser.parse_args()

    main(args.title, args.output_docx, args.screenshot_folder)
